"""报告 docx 构建助手（report-writing skill 配套）。

定位：把已写好的报告内容渲染成符合 skill §四 规范的 docx——
中文宋体 / 英文 Times New Roman、三线表、表上图下题注、纯黑、统计符号斜体、
干净的居中加粗标题页（无灰色说明小字）。

用法（在生成脚本里 import）：
    from build_report import Report
    rep = Report()
    rep.title_lines(["一项 ... 研究", "24--36 周体重反弹补充分析报告"])  # 多行居中加粗，无副标题灰字
    rep.heading("一、分析背景与目的", level=1)
    rep.para("本补充分析用于评估 ...")                 # 完整段落；statt 符号见 rep.para_runs
    rep.table_caption("表1 24--36周体重反弹分析样本量")
    rep.three_line_table(header=[...], rows=[...])      # 或 rep.table_from_xlsx(path, sheet)
    rep.note("注：随机入组 N 按 ...")
    rep.figure("../04_figures/Fig1_xxx.png", caption="图1 各组体重变化轨迹")
    rep.save("报告_v1.docx", also_md=True)              # 默认同时落 .md

正文内容必须由调用方按 skill 铁律手写（数据有源、完整段落、零编造），
本模块只负责"排版正确"，不负责"内容生成"。
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CN_BODY = "宋体"           # 中文正文
CN_HEAD = "黑体"           # 中文标题（可换"微软雅黑"）
EN_FONT = "Times New Roman"  # 英文/数字
BLACK = (0, 0, 0)


def val(yaml_path, key, which="full"):
    """从结果单一真源 results.yaml 取已渲染成品串（C1：下游取数禁手敲）。
    数字只在 results.yaml 改；改下游须先回写真源再传播。
    用法：rep.para("S2 vs S1 差异为 " + val("07_paper/results.yaml", "S2_vs_S1_diff") + "。")
    """
    import yaml
    with open(yaml_path, "r", encoding="utf-8") as f:
        doc = yaml.safe_load(f) or {}
    res = (doc.get("results") or {}).get(key)
    if res is None:
        raise KeyError(f"results.yaml 无键：{key}")
    s = (res.get("rendered") or {}).get(which)
    if s is None:
        raise KeyError(f"键 {key} 无 rendered.{which}")
    return s


def setfont(run, cn=CN_BODY, en=EN_FONT, size=10.5, bold=False, italic=False, color=BLACK):
    """每个 run 都同时设英文字体与中文 eastAsia 字体，否则字体会回退。"""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = en
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), en)
    rfonts.set(qn("w:hAnsi"), en)
    rfonts.set(qn("w:eastAsia"), cn)
    run.font.color.rgb = RGBColor(*color)


def _set_cell_border(cell, **edges):
    """给单元格设指定边框（用于三线表：只在需要的行设顶/底线）。"""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge in ("top", "bottom", "left", "right"):
        spec = edges.get(edge)
        tag = qn(f"w:{edge}")
        el = borders.find(tag)
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        if spec:  # spec = (sz_eighths_pt,)
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(spec[0]))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "none")


class Report:
    def __init__(self, body_size=10.5):
        self.doc = Document()
        self.body_size = body_size
        self._md = []  # 同步累积 markdown
        # 默认正文样式字体（兜底；真正生效靠每个 run setfont）
        normal = self.doc.styles["Normal"]
        normal.font.name = EN_FONT
        normal.font.size = Pt(body_size)
        normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CN_BODY)

    # ---- 标题页 ----
    def title_lines(self, lines, size=16):
        """多行居中加粗标题；不加任何灰色说明小字。"""
        for ln in lines:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)
            setfont(p.add_run(ln), cn=CN_HEAD, size=size, bold=True)
        self.doc.add_paragraph()
        self._md.append("# " + " ".join(lines) + "\n")

    def meta(self, text):
        """可选的日期/版本，纯黑正常字号、居中。绝不灰、不缩小、不加解说。"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        setfont(p.add_run(text), size=self.body_size)
        self._md.append(f"*{text}*\n")

    # ---- 章节 ----
    def heading(self, text, level=1):
        size = {1: 15, 2: 13, 3: 12}.get(level, 12)
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        setfont(p.add_run(text), cn=CN_HEAD, size=size, bold=True)
        self._md.append("\n" + "#" * (level + 1) + " " + text + "\n")

    # ---- 段落 ----
    def para(self, text, size=None):
        """普通完整段落。"""
        p = self.doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(4)
        setfont(p.add_run(text), size=size or self.body_size)
        self._md.append(text + "\n")

    def para_runs(self, runs, size=None):
        """混合排版段落：runs = [(text, {italic:True}), ...]，用于把 P、vs 等设斜体。"""
        p = self.doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(4)
        md = []
        for text, opts in runs:
            opts = opts or {}
            setfont(p.add_run(text), size=size or self.body_size,
                    italic=opts.get("italic", False), bold=opts.get("bold", False))
            md.append(f"*{text}*" if opts.get("italic") else text)
        self._md.append("".join(md) + "\n")

    def summary_item(self, label, text):
        """执行摘要项：加粗标签 + 整句，不用项目符号点。"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        setfont(p.add_run(label + "："), size=self.body_size, bold=True)
        setfont(p.add_run(text), size=self.body_size)
        self._md.append(f"**{label}：**{text}\n")

    # ---- 表 ----
    def table_caption(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        setfont(p.add_run(text), size=self.body_size, bold=True)
        self._md.append("\n**" + text + "**\n")

    def note(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        setfont(p.add_run(text), size=self.body_size - 1.5, italic=True)
        self._md.append(f"*{text}*\n")

    def three_line_table(self, header, rows, num_cols_right=None):
        """三线表：顶线/表头下线/底线，无竖线无内部横线。
        num_cols_right: 右对齐的列索引集合（数字列）；默认除第 1 列外全右对齐。
        """
        ncol = len(header)
        if num_cols_right is None:
            num_cols_right = set(range(1, ncol))
        t = self.doc.add_table(rows=1, cols=ncol)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        big, small = 12, 4  # 边框粗细（1/8 pt）
        # 表头
        for j, htext in enumerate(header):
            cell = t.rows[0].cells[j]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            setfont(p.add_run(str(htext)), size=self.body_size, bold=True)
            _set_cell_border(cell, top=(big,), bottom=(small,))
        # 数据行
        for i, row in enumerate(rows):
            cells = t.add_row().cells
            last = (i == len(rows) - 1)
            for j, val in enumerate(row):
                cell = cells[j]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p = cell.paragraphs[0]
                p.alignment = (WD_ALIGN_PARAGRAPH.RIGHT if j in num_cols_right
                               else WD_ALIGN_PARAGRAPH.LEFT)
                setfont(p.add_run("" if val is None else str(val)), size=self.body_size)
                _set_cell_border(cell, bottom=(big,) if last else None)
        # markdown 镜像
        self._md.append("| " + " | ".join(map(str, header)) + " |")
        self._md.append("| " + " | ".join(["---"] * ncol) + " |")
        for row in rows:
            self._md.append("| " + " | ".join("" if v is None else str(v) for v in row) + " |")
        self._md.append("")
        return t

    def table_from_xlsx(self, path, sheet=0, header_row=0, max_rows=None):
        """从 03_tables/ 的 xlsx 读一张表填三线表（自动取数，不手敲）。"""
        from openpyxl import load_workbook
        wb = load_workbook(path, data_only=True)
        ws = wb.worksheets[sheet] if isinstance(sheet, int) else wb[sheet]
        data = [[c if c is not None else "" for c in r]
                for r in ws.iter_rows(values_only=True)]
        data = [r for r in data if any(str(c).strip() for c in r)]  # 去空行
        header = data[header_row]
        rows = data[header_row + 1:]
        if max_rows:
            rows = rows[:max_rows]
        return self.three_line_table(header, rows)

    # ---- 图 ----
    def figure(self, path, caption=None, width_in=6.0):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(width_in))
        if caption:
            c = self.doc.add_paragraph()
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.paragraph_format.space_before = Pt(2)
            setfont(c.add_run(caption), size=self.body_size, bold=True)
        self._md.append(f"\n![{caption or ''}]({path})")
        if caption:
            self._md.append(f"**{caption}**\n")

    # ---- 落盘 ----
    def save(self, docx_path, also_md=True):
        self.doc.save(docx_path)
        out = [docx_path]
        if also_md:
            md_path = docx_path.rsplit(".", 1)[0] + ".md"
            with open(md_path, "w", encoding="utf-8") as f:
                f.write("\n".join(self._md).rstrip() + "\n")
            out.append(md_path)
        return out


if __name__ == "__main__":  # 冒烟自测
    rep = Report()
    rep.title_lines(["示例：某随机对照研究", "24--36 周体重反弹补充分析报告"])
    rep.meta("2026-06-14 ｜ v1")
    rep.heading("一、分析背景与目的", level=1)
    rep.para("本补充分析用于评估第 24 周停止干预后至第 36 周期间三组受试者的体重回升情况。")
    rep.para_runs([("主要研究方向 S2 vs S1 的组间反弹差异为 −1.82 kg（95%CI：−3.29，−0.36），",
                    None), ("P", {"italic": True}), (" = 0.015，差异达统计学显著。", None)])
    rep.table_caption("表1 各组样本量")
    rep.three_line_table(
        header=["组别编码", "组别", "随机入组 N", "24/36周例数"],
        rows=[["S0", "S0（单用）", 27, 22], ["S1", "S1（对照）", 28, 23], ["S2", "S2（联合）", 26, 24]],
    )
    rep.note("注：以同时具有第 24 周和第 36 周记录者作为完全病例集。")
    paths = rep.save("_smoketest_report.docx", also_md=True)
    print("OK:", paths)
